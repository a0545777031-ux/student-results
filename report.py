# -*- coding: utf-8 -*-
"""Build graphical reports (PDF via reportlab, Word via python-docx) from a
payload of chart images + student ratings. Arabic-aware."""
import io, base64, os
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
import ar_shape as S
import data_io

LEVELS = [("excellent","ممتاز","1e7d4f"), ("vgood","جيد جداً","2f8f79"),
          ("good","جيد","b6892b"), ("pass","مقبول","d08c3a"), ("weak","ضعيف","b02a37")]
LEVEL_AR = {k:ar for k,ar,_ in LEVELS}

def _img_bytes(dataurl):
    if not dataurl: return None
    if "," in dataurl: dataurl = dataurl.split(",",1)[1]
    try: return io.BytesIO(base64.b64decode(dataurl))
    except Exception: return None

# ---------------- PDF ----------------
def build_pdf(payload, out_path):
    data_io._ensure_fonts()
    W,H = A4
    c = canvas.Canvas(out_path, pagesize=A4)
    M = 40
    def rtl(t): return S.shape(str(t))
    y = [H-50]
    def header():
        c.setFillColorRGB(0.055,0.353,0.302); c.rect(0,H-70,W,70,fill=1,stroke=0)
        c.setFillColorRGB(1,1,1); c.setFont("Ar-Bold",17)
        c.drawRightString(W-M,H-38,rtl("تقرير تحليل نتائج الطلاب"))
        c.setFont("Ar",10)
        meta=payload.get("meta",{})
        sub=" | ".join([x for x in [meta.get("school",""),meta.get("grade_class",""),meta.get("year","")] if x])
        c.drawRightString(W-M,H-57,rtl(sub))
        c.setFillColorRGB(0.78,0.64,0.29); c.setFont("Ar-Bold",12)
        c.drawString(M,H-40,rtl(payload.get("title","")))
        y[0]=H-95
    def need(space):
        if y[0]-space < 50:
            c.showPage(); header()
    header()
    # ratings section
    c.setFillColorRGB(0.06,0.30,0.26); c.setFont("Ar-Bold",14)
    c.drawRightString(W-M,y[0],rtl("تقديرات المواد")); y[0]-=8
    c.setStrokeColorRGB(0.85,0.9,0.88); c.line(M,y[0],W-M,y[0]); y[0]-=20
    for k,ar,hexc in LEVELS:
        names = (payload.get("classification",{}).get(k) or [])
        need(40)
        r=int(hexc[0:2],16)/255; g=int(hexc[2:4],16)/255; b=int(hexc[4:6],16)/255
        c.setFillColorRGB(r,g,b); c.roundRect(W-M-150,y[0]-4,150,20,4,fill=1,stroke=0)
        c.setFillColorRGB(1,1,1); c.setFont("Ar-Bold",11)
        c.drawCentredString(W-M-75,y[0]+1,rtl(f"{ar} ({len(names)})"))
        y[0]-=24
        c.setFillColorRGB(0.15,0.15,0.15); c.setFont("Ar",10)
        if names:
            # wrap names across the width
            line=""; maxw=W-2*M
            for nm in names:
                cand=(line+"،  "+nm) if line else nm
                if S.shape(cand) and c.stringWidth(rtl(cand),"Ar",10) > maxw:
                    need(16); c.drawRightString(W-M,y[0],rtl(line)); y[0]-=15; line=nm
                else: line=cand
            if line: need(16); c.drawRightString(W-M,y[0],rtl(line)); y[0]-=15
        else:
            need(16); c.setFillColorRGB(0.5,0.5,0.5); c.drawRightString(W-M,y[0],rtl("لا يوجد")); y[0]-=15
        y[0]-=8
    # charts section
    imgs = payload.get("images",[])
    if imgs:
        need(30); c.setFillColorRGB(0.06,0.30,0.26); c.setFont("Ar-Bold",14)
        c.drawRightString(W-M,y[0],rtl("الرسوم البيانية")); y[0]-=8
        c.setStrokeColorRGB(0.85,0.9,0.88); c.line(M,y[0],W-M,y[0]); y[0]-=14
        for im in imgs:
            b=_img_bytes(im.get("data"))
            if not b: continue
            try: ir=ImageReader(b); iw,ih=ir.getSize()
            except Exception: continue
            dw=W-2*M; dh=dw*ih/iw
            if dh>300: dh=300; dw=dh*iw/ih
            need(dh+26)
            c.setFillColorRGB(0.1,0.1,0.1); c.setFont("Ar-Bold",11)
            c.drawRightString(W-M,y[0],rtl(im.get("title",""))); y[0]-=dh+6
            c.drawImage(ir,(W-dw)/2,y[0],dw,dh,preserveAspectRatio=True,mask='auto'); y[0]-=16
    c.showPage(); c.save(); return out_path

# ---------------- Word ----------------
def build_docx(payload, out_path):
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    # default font + RTL for normal style
    style = doc.styles["Normal"]; style.font.name="Arial"; style.font.size=Pt(11)
    def set_rtl(p):
        pPr=p._p.get_or_add_pPr(); bidi=OxmlElement('w:bidi'); pPr.append(bidi)
        p.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    def heading(text,size=15,color="0E5A4D"):
        p=doc.add_paragraph(); set_rtl(p); r=p.add_run(text); r.bold=True
        r.font.size=Pt(size); r.font.color.rgb=RGBColor.from_string(color); return p

    meta=payload.get("meta",{})
    heading("تقرير تحليل نتائج الطلاب",18)
    p=doc.add_paragraph(); set_rtl(p); r=p.add_run(payload.get("title","")); r.bold=True
    r.font.size=Pt(13); r.font.color.rgb=RGBColor.from_string("B6892B")
    sub=" | ".join([x for x in [meta.get("school",""),meta.get("grade_class",""),meta.get("year","")] if x])
    if sub: p=doc.add_paragraph(); set_rtl(p); p.add_run(sub).font.size=Pt(10)

    heading("تقديرات المواد",15)
    for k,ar,hexc in LEVELS:
        names=(payload.get("classification",{}).get(k) or [])
        p=doc.add_paragraph(); set_rtl(p)
        r=p.add_run(f"{ar}: {len(names)} طالب"); r.bold=True; r.font.size=Pt(12)
        r.font.color.rgb=RGBColor.from_string(hexc.upper())
        p2=doc.add_paragraph(); set_rtl(p2)
        p2.add_run("، ".join(names) if names else "لا يوجد").font.size=Pt(10)

    imgs=payload.get("images",[])
    if imgs:
        heading("الرسوم البيانية",15)
        for im in imgs:
            b=_img_bytes(im.get("data"))
            if not b: continue
            p=doc.add_paragraph(); set_rtl(p); r=p.add_run(im.get("title","")); r.bold=True; r.font.size=Pt(11)
            try:
                pic=doc.add_paragraph(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER
                pic.add_run().add_picture(b, width=Inches(6.2))
            except Exception:
                pass
    doc.save(out_path); return out_path

